# src/application/services/documento_export_service.py
"""
Servicio para la exportación de documentos a Word.
"""
import os
import json
import datetime
from typing import Dict, Any, Optional, List, BinaryIO
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from io import BytesIO
from PIL import Image

from domain.models.documento import Documento
from domain.models.nodo_ipran import NodoIPRAN
from application.services.documento_service import DocumentoService
from application.services.nodo_ipran_service import NodoIPRANService

class DocumentoExportService:
    """Servicio para exportar documentos a formato Word."""
    
    def __init__(self):
        """Constructor del servicio."""
        self.documento_service = DocumentoService()
        self.nodo_service = NodoIPRANService()
        # Directorio de recursos para plantillas y almacenamiento de documentos
        self.recursos_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "recursos")
        self.plantillas_dir = os.path.join(self.recursos_dir, "plantillas")
        self.docs_dir = os.path.join(self.recursos_dir, "documentos")
        # Crear directorios si no existen
        os.makedirs(self.plantillas_dir, exist_ok=True)
        os.makedirs(self.docs_dir, exist_ok=True)
    
    def exportar_a_word(self, documento_id: int) -> str:
        """
        Exporta un documento a formato Word.
        
        Args:
            documento_id: ID del documento a exportar
            
        Returns:
            str: Ruta del archivo Word generado
            
        Raises:
            ValueError: Si el documento no existe o si faltan datos necesarios
        """
        # Obtener el documento
        documento = self.documento_service.obtener_por_id(documento_id)
        if not documento:
            raise ValueError(f"No existe un documento con ID {documento_id}")
        
        # Obtener datos adicionales necesarios
        nodo = None
        if documento.nodo_id:
            nodo = self.nodo_service.obtener_por_id(documento.nodo_id)
        
        # Cargar el contenido JSON
        contenido = {}
        if documento.contenido_json:
            contenido = json.loads(documento.contenido_json)
        
        # Crear un nuevo documento Word
        doc = DocxDocument()
        
        # Configurar el documento según la topología
        if documento.tipo_topologia == "IPRAN+MIKROTIK":
            self._generar_documento_ipran_mikrotik(doc, documento, nodo, contenido)
        elif documento.tipo_topologia == "IPRAN+RADWIN":
            self._generar_documento_ipran_radwin(doc, documento, nodo, contenido)
        elif documento.tipo_topologia == "IPRAN+GPON":
            self._generar_documento_ipran_gpon(doc, documento, nodo, contenido)
        else:
            # Topología genérica
            self._generar_documento_generico(doc, documento, nodo, contenido)
        
        # Generar el nombre del archivo
        fecha_str = documento.fecha_creacion.strftime("%Y%m%d")
        nombre_archivo = f"{fecha_str}_{documento.tipo_transaccion}_{documento.cliente_id}.docx"
        ruta_archivo = os.path.join(self.docs_dir, nombre_archivo)
        
        # Guardar el documento
        doc.save(ruta_archivo)
        
        return ruta_archivo
    
    def _generar_documento_ipran_mikrotik(self, doc: DocxDocument, documento: Documento, 
                                          nodo: Optional[NodoIPRAN], contenido: Dict[str, Any]) -> None:
        """
        Genera un documento Word para la topología IPRAN+MIKROTIK.
        
        Args:
            doc: Documento Word a generar
            documento: Información del documento
            nodo: Información del nodo IPRAN
            contenido: Contenido adicional del documento
        """
        # Configurar estilos
        self._configurar_estilos_documento(doc)
        
        # Agregar encabezado
        self._agregar_encabezado(doc, documento)
        
        # Etiqueta del cliente
        etiqueta = self.documento_service.generar_etiqueta_cliente(
            documento.cliente_id, documento.cliente_nombre, documento.ancho_banda)
        doc.add_paragraph(etiqueta)
        
        # Agregar sección de configuración de IPRAN
        self._agregar_seccion_titulo(doc, "CONFIGURACION NODO IPRAN")
        
        if nodo:
            # Tabla de información del nodo
            table = doc.add_table(rows=7, cols=2)
            table.style = 'Table Grid'
            
            # Configurar las celdas
            self._agregar_fila_tabla(table, 0, "NODO:", nodo.nombre_nodo)
            self._agregar_fila_tabla(table, 1, "SWITCH:", nodo.alias_nodo)
            self._agregar_fila_tabla(table, 2, "IP SWITCH:", contenido.get("ip_switch", ""))
            self._agregar_fila_tabla(table, 3, "PUERTO:", contenido.get("puerto", ""))
            self._agregar_fila_tabla(table, 4, "VLAN:", contenido.get("vlan", ""))
            self._agregar_fila_tabla(table, 5, "IP MIKROTIK:", documento.mikrotik_ip or "")
            self._agregar_fila_tabla(table, 6, "IP PUBLICA:", contenido.get("ip_publica", ""))
        
        # Agregar sección de configuración de Mikrotik
        self._agregar_seccion_titulo(doc, "CONFIGURACION MIKROTIK")
        
        # Agregar export de Mikrotik si existe
        if "mikrotik_export" in contenido:
            p = doc.add_paragraph(contenido["mikrotik_export"])
            p.style = 'Code'
        
        # Agregar sección de gráfica de consumo
        self._agregar_seccion_titulo(doc, "GRAFICA DE CONSUMO POR EL CLIENTE")
        
        # Agregar imagen de gráfica si existe
        if "grafica_consumo" in contenido:
            self._agregar_imagen_desde_bytes(doc, contenido["grafica_consumo"])
        
        # Agregar sección de link solarwinds
        if "link_solarwinds" in contenido:
            p = doc.add_paragraph("LINK SOLARWINDS")
            p.style = 'Heading 2'
            doc.add_paragraph(contenido["link_solarwinds"])
        
        # Agregar sección de notas/observaciones
        self._agregar_seccion_titulo(doc, "NOTAS/OBSERVACIONES")
        
        if "observaciones" in contenido:
            for obs in contenido["observaciones"]:
                p = doc.add_paragraph(obs, style='List Bullet')
        
        # Agregar sección de correo de notificación
        self._agregar_seccion_titulo(doc, "CORREO DE NOTIFICACION")
        
        if "correo" in contenido:
            doc.add_paragraph(contenido["correo"])
    
    def _generar_documento_ipran_radwin(self, doc: DocxDocument, documento: Documento, 
                                       nodo: Optional[NodoIPRAN], contenido: Dict[str, Any]) -> None:
        """
        Genera un documento Word para la topología IPRAN+RADWIN.
        
        Args:
            doc: Documento Word a generar
            documento: Información del documento
            nodo: Información del nodo IPRAN
            contenido: Contenido adicional del documento
        """
        # Implementación similar a IPRAN+MIKROTIK pero con secciones específicas para RADWIN
        # Por simplicidad, esto puede ser extendido más tarde según necesidades específicas
        self._generar_documento_generico(doc, documento, nodo, contenido)
    
    def _generar_documento_ipran_gpon(self, doc: DocxDocument, documento: Documento, 
                                     nodo: Optional[NodoIPRAN], contenido: Dict[str, Any]) -> None:
        """
        Genera un documento Word para la topología IPRAN+GPON.
        
        Args:
            doc: Documento Word a generar
            documento: Información del documento
            nodo: Información del nodo IPRAN
            contenido: Contenido adicional del documento
        """
        # Implementación similar a IPRAN+MIKROTIK pero con secciones específicas para GPON
        # Por simplicidad, esto puede ser extendido más tarde según necesidades específicas
        self._generar_documento_generico(doc, documento, nodo, contenido)
    
    def _generar_documento_generico(self, doc: DocxDocument, documento: Documento, 
                                   nodo: Optional[NodoIPRAN], contenido: Dict[str, Any]) -> None:
        """
        Genera un documento Word genérico con la información básica.
        
        Args:
            doc: Documento Word a generar
            documento: Información del documento
            nodo: Información del nodo IPRAN
            contenido: Contenido adicional del documento
        """
        # Configurar estilos
        self._configurar_estilos_documento(doc)
        
        # Agregar encabezado
        self._agregar_encabezado(doc, documento)
        
        # Etiqueta del cliente
        etiqueta = self.documento_service.generar_etiqueta_cliente(
            documento.cliente_id, documento.cliente_nombre, documento.ancho_banda)
        doc.add_paragraph(etiqueta)
        
        # Recorrer todas las secciones en el contenido y agregarlas al documento
        for seccion, datos in contenido.items():
            if seccion == "imagenes" and isinstance(datos, dict):
                # Procesar imágenes
                for titulo_imagen, imagen_bytes in datos.items():
                    self._agregar_seccion_titulo(doc, titulo_imagen)
                    self._agregar_imagen_desde_bytes(doc, imagen_bytes)
            elif seccion == "textos" and isinstance(datos, dict):
                # Procesar textos
                for titulo_texto, texto in datos.items():
                    self._agregar_seccion_titulo(doc, titulo_texto)
                    doc.add_paragraph(texto)
            elif seccion == "observaciones" and isinstance(datos, list):
                # Procesar observaciones
                self._agregar_seccion_titulo(doc, "NOTAS/OBSERVACIONES")
                for obs in datos:
                    p = doc.add_paragraph(obs, style='List Bullet')
    
    def _configurar_estilos_documento(self, doc: DocxDocument) -> None:
        """
        Configura los estilos del documento Word.
        
        Args:
            doc: Documento Word a configurar
        """
        # Configurar estilo para títulos
        style = doc.styles['Heading 1']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = (255, 0, 0)  # Rojo
        
        style = doc.styles['Heading 2']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = True
        font.color.rgb = (255, 0, 0)  # Rojo
        
        # Configurar estilo para código
        style = doc.styles.add_style('Code', 1)
        font = style.font
        font.name = 'Consolas'
        font.size = Pt(9)
    
    def _agregar_encabezado(self, doc: DocxDocument, documento: Documento) -> None:
        """
        Agrega el encabezado al documento Word.
        
        Args:
            doc: Documento Word
            documento: Información del documento
        """
        # Crear tabla de encabezado (3 columnas)
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        
        # Logo (celda 1)
        # Aquí se podría agregar un logo si se tiene disponible
        
        # Título (celda 2)
        cell = table.cell(0, 1)
        title_paragraph = cell.paragraphs[0]
        title_paragraph.alignment = 1  # Centrado
        title_run = title_paragraph.add_run(f"{documento.tipo_transaccion} DE ENLACE DE INTERNET\n{documento.cliente_nombre}")
        title_run.bold = True
        title_run.font.size = Pt(12)
        
        # Versión y página (celda 3)
        cell = table.cell(0, 2)
        version_paragraph = cell.paragraphs[0]
        version_paragraph.alignment = 1  # Centrado
        version_run = version_paragraph.add_run("Ver:\t1.0\nPágina\t1 de 1")
        version_run.font.size = Pt(10)
        
        # Fecha e ingeniero (celda inferior)
        fecha_str = documento.fecha_creacion.strftime("%d %b %Y")
        cell = table.cell(1, 0)
        cell.merge(table.cell(1, 2))
        fecha_paragraph = cell.paragraphs[0]
        fecha_run = fecha_paragraph.add_run(f"FECHA: {fecha_str}\nTX Access Engineer: {documento.ingeniero}")
        fecha_run.font.size = Pt(10)
    
    def _agregar_seccion_titulo(self, doc: DocxDocument, titulo: str) -> None:
        """
        Agrega un título de sección al documento.
        
        Args:
            doc: Documento Word
            titulo: Título de la sección
        """
        doc.add_paragraph(titulo, style='Heading 2')
        # Línea separadora
        p = doc.add_paragraph()
        p.add_run("*" * 70)
    
    def _agregar_fila_tabla(self, table, row_index: int, label: str, value: str) -> None:
        """
        Agrega una fila a una tabla con etiqueta y valor.
        
        Args:
            table: Tabla a la que agregar la fila
            row_index: Índice de la fila
            label: Etiqueta (primera columna)
            value: Valor (segunda columna)
        """
        cell_label = table.cell(row_index, 0)
        cell_label.text = label
        cell_label.paragraphs[0].runs[0].bold = True
        
        cell_value = table.cell(row_index, 1)
        cell_value.text = value
    
    def _agregar_imagen_desde_bytes(self, doc: DocxDocument, imagen_bytes: bytes) -> None:
        """
        Agrega una imagen al documento a partir de bytes.
        
        Args:
            doc: Documento Word
            imagen_bytes: Bytes de la imagen
        """
        # Crear un stream temporal para la imagen
        stream = BytesIO(imagen_bytes)
        # Abrir la imagen con PIL para obtener el tamaño
        with Image.open(stream) as img:
            width, height = img.size
        
        # Reiniciar el stream
        stream.seek(0)
        
        # Calcular el ancho máximo para la imagen (ajustando a la página)
        max_width = Inches(6)  # Ancho máximo de 6 pulgadas
        
        # Calcular la relación de aspecto
        if width > max_width.pt:
            ratio = max_width.pt / width
            width = max_width.pt
            height = height * ratio
        
        # Agregar la imagen al documento
        doc.add_picture(stream, width=Pt(width), height=Pt(height))
        
    def guardar_imagen(self, imagen_bytes: bytes, nombre: str) -> str:
        """
        Guarda una imagen en el sistema de archivos.
        
        Args:
            imagen_bytes: Bytes de la imagen
            nombre: Nombre base para el archivo
            
        Returns:
            str: Ruta del archivo guardado
        """
        # Crear un directorio para imágenes si no existe
        imagenes_dir = os.path.join(self.docs_dir, "imagenes")
        os.makedirs(imagenes_dir, exist_ok=True)
        
        # Generar un nombre único
        nombre_archivo = f"{nombre}_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.png"
        ruta_archivo = os.path.join(imagenes_dir, nombre_archivo)
        
        # Guardar la imagen
        with open(ruta_archivo, 'wb') as f:
            f.write(imagen_bytes)
        
        return ruta_archivo